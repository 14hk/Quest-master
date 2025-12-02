import os
from jinja2 import Environment, FileSystemLoader
from docx import Document
from datetime import datetime

try:
    from weasyprint import HTML
    _WEASY_AVAILABLE = True
except OSError:
    _WEASY_AVAILABLE = False
    print("WeasyPrint (GTK) не установлен, экспорт в PDF недоступен.")

TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), '..', 'templates')
PARCHMENTS_DIR = os.path.join(os.path.dirname(__file__), '..', '..', 'parchments')
if not os.path.exists(PARCHMENTS_DIR):
    os.makedirs(PARCHMENTS_DIR)

env = Environment(loader=FileSystemLoader(TEMPLATE_DIR))

def render_template(template_name, context):
    template = env.get_template(template_name)
    return template.render(context)

def export_pdf(quest_data, template_name="guild_contract.html"):
    if not _WEASY_AVAILABLE:
        raise RuntimeError("WeasyPrint (GTK) не установлен, экспорт в PDF недоступен.")

    context = {"quest": quest_data, "timestamp": datetime.now()}
    html_content = render_template(template_name, context)
    
    filename = f"{quest_data['id']}_{int(datetime.now().timestamp())}.pdf"
    filepath = os.path.join(PARCHMENTS_DIR, filename)
    
    HTML(string=html_content).write_pdf(filepath)
    return filepath

def export_docx(quest_data):
    doc = Document()
    doc.add_heading(f"Квест #{quest_data['id']}: {quest_data['title']}", 0)
    
    doc.add_paragraph(f"Сложность: {quest_data['difficulty']}")
    doc.add_paragraph(f"Награда: {quest_data['reward']} золотых")
    doc.add_paragraph(f"Дедлайн: {quest_data['deadline']}")
    doc.add_heading('Описание', level=1)
    doc.add_paragraph(quest_data['description'])
    
    filename = f"{quest_data['id']}_{int(datetime.now().timestamp())}.docx"
    filepath = os.path.join(PARCHMENTS_DIR, filename)
    doc.save(filepath)
    return filepath

class BatchExporter:
    @staticmethod
    def generate_100_quests():
        
        dummy_quest = {
            "id": 999,
            "title": "Speedrun po pdf, pognali",
            "difficulty": "Эпический",
            "reward": 1000,
            "description": "Сверхзвуковая генерация квестов",
            "deadline": "Now"
        }
        
        for i in range(100):
            dummy_quest["id"] = i
            html = render_template("guild_contract.html", {"quest": dummy_quest})
            fname = os.path.join(PARCHMENTS_DIR, f"batch_{i}.html")
            with open(fname, "w", encoding="utf-8") as f:
                f.write(html)
